import itertools
import os
import platform
import re
import time
from concurrent.futures import ThreadPoolExecutor

import pandas as pd
from docx import Document

if platform.system() == "Windows":
    import pythoncom
    from docx2pdf import convert


def docx_replace_regex(doc_obj: Document, regex_to_replace: re.compile, replacement: str) -> None:
    """
    Replace the regex in a docx.Document object.

    Args:
        doc_obj: Template `docx.Document` object.
        regex_to_replace: Regex compile to replace in the document.
        replacement : Replacement string to substitute for the regex.
    """
    for p in doc_obj.paragraphs:
        if regex_to_replace.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex_to_replace.search(inline[i].text):
                    text = regex_to_replace.sub(replacement, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex_to_replace, replacement)


def filename_ext(filename: str) -> str:
    return os.path.splitext(filename)[1]


def generate_doc(
    doc_spec: tuple[int, pd.Series],
    template_docx: str,
    output_dir: str,
    output_base_fn: str,
    output_filetype: str,
) -> str:
    doc_number = doc_spec[0]
    doc_replacements = doc_spec[1]

    if platform.system() == "Windows":
        pythoncom.CoInitialize()  # needed to prevent docx2pdf multithreading error

    _log = []

    _log.append(f"Document {doc_number+1} - creating replacements for:")
    doc = Document(template_docx)
    output_fn_additions = []

    # Loop through each replacement in the row for the document
    for str_to_replace, replacement_str in doc_replacements.items():
        _log.append(f"\t{str_to_replace} -> {replacement_str}")

        # Add column values that are not excluded to list to be added to
        # output filename
        column_substrings_excl_from_fn = [
            "date",
            "industry",
        ]
        if not any(substr in str_to_replace.lower() for substr in column_substrings_excl_from_fn):
            output_fn_additions.append(replacement_str)

        # Execute document replacement
        regex = re.compile(re.escape(str_to_replace))
        docx_replace_regex(doc, regex, replacement_str)

    # Save .docx file
    output_fn_addition_str = " - ".join(output_fn_additions)
    if output_fn_addition_str:
        output_fn_addition_str = f" - {output_fn_addition_str}"
    output_fn = os.path.join(output_dir, f"{output_base_fn}{output_fn_addition_str}.docx")
    doc.save(output_fn)

    if output_filetype == ".pdf":
        # Convert .docx file to .pdf file
        convert(output_fn)
        os.remove(output_fn)  # remove .docx version
        output_fn = output_fn.replace(".docx", ".pdf")

    _log.append(f"Document {doc_number+1} - file saved to '{output_fn}'.")

    return output_fn


def batch_replace(
    *,
    template_docx: str,
    replacements_csv: str,
    max_new_docs: int = 25,
    output_dir: str,
    output_base_fn: str,
    output_filetype: str,
) -> list[str]:
    """
    Generate multiple output files by executing multiple sets of find-replace operations.

    Args:
        template_docx: Path to the template document.
        replacements_csv: Path to the replacements spec document. Each row
            specifies a document version to generate. Each column header
            specifies the text to replace in the template. The column values
            specify the text to substitute during the replacement.
        output_dir: Path to the output directory.
        output_base_fn: Base filename for the generated output files.
        output_filetype: Output filetype (.docx, .pdf)
    """

    replacements_df = pd.read_csv(replacements_csv)
    replacements_df = replacements_df.truncate(after=max_new_docs - 1)  # limit number of documents generated
    output_filetype = output_filetype if output_filetype in [".docx", ".pdf"] else ".pdf"

    # Check input file extensions
    if filename_ext(template_docx) != ".docx":
        raise ValueError(f"{template_docx} does not have a valid file extension.")
    if filename_ext(replacements_csv) != ".csv":
        raise ValueError(f"{replacements_csv} does not have a valid file extension.")

    start = time.time()

    # # Loop through each replacement row
    # # doc_spec is containerized in tuple to match ThreadPoolExecutor behavior
    # output_file_paths = []
    # for doc_number, doc_replacements in replacements_df.iterrows():
    #     output_file_paths.append(
    #         generate_doc(
    #             (doc_number, doc_replacements),
    #             template_docx,
    #             output_dir,
    #             output_base_fn,
    #             output_filetype,
    #         )
    #     )

    with ThreadPoolExecutor() as executor:
        output_file_paths = executor.map(
            generate_doc,
            replacements_df.iterrows(),
            itertools.repeat(template_docx),
            itertools.repeat(output_dir),
            itertools.repeat(output_base_fn),
            itertools.repeat(output_filetype),
        )

    print("Time elapsed:", time.time() - start)

    return output_file_paths


if __name__ == "__main__":
    output_file_paths = batch_replace(
        template_docx="cover_letter_test.docx",
        replacements_csv="replacements.csv",
        max_new_docs=5,
        output_dir="created",
        output_base_fn="Cover Letter Test",
        output_filetype=".docx",
    )

    from os.path import basename
    from zipfile import ZipFile

    with ZipFile("generated_documents.zip", "w") as zip_obj:
        for generated_doc_path in output_file_paths:
            zip_obj.write(generated_doc_path, basename(generated_doc_path))